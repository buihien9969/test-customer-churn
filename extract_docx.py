import docx
import os

# Đường dẫn đến file Word
file_path = os.path.join(os.getcwd(), "20231102_EntranceTest_DAInternMCNA.docx")

# Đọc file Word
doc = docx.Document(file_path)

# In ra nội dung của file
print("=== NỘI DUNG CỦA FILE WORD ===")
for para in doc.paragraphs:
    if para.text.strip():  # Chỉ in ra các đoạn không trống
        print(para.text)
        print("---")  # Phân cách giữa các đoạn

# In ra nội dung của các bảng (nếu có)
print("\n=== NỘI DUNG CỦA CÁC BẢNG ===")
for i, table in enumerate(doc.tables):
    print(f"Bảng {i+1}:")
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(" | ".join(row_text))
    print("---")
